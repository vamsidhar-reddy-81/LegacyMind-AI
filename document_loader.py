from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

FAST_TEXT_CHARS = 45_000
FAST_PDF_PAGES = 10
DEEP_TEXT_CHARS = 180_000
DEEP_PDF_PAGES = 80


@dataclass
class LoadedDocument:
    name: str
    text: str


def load_file(name: str, data: bytes, fast_mode: bool = True) -> LoadedDocument:
    suffix = Path(name).suffix.lower()
    if suffix == ".txt":
        text = data.decode("utf-8", errors="ignore")
    elif suffix == ".pdf":
        text = _load_pdf(data, fast_mode=fast_mode)
    elif suffix == ".docx":
        text = _load_docx(data)
    else:
        raise ValueError(f"Unsupported file type: {suffix or 'unknown'}")

    max_chars = FAST_TEXT_CHARS if fast_mode else DEEP_TEXT_CHARS
    return LoadedDocument(name=name, text=clean_text(text)[:max_chars])


def clean_text(text: str) -> str:
    return " ".join(text.replace("\x00", " ").split())


def _load_pdf(data: bytes, fast_mode: bool) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("Install pypdf to read PDF files.") from exc

    reader = PdfReader(BytesIO(data))
    max_pages = FAST_PDF_PAGES if fast_mode else DEEP_PDF_PAGES
    pages = reader.pages[:max_pages]
    return "\n".join(_clean_pdf_page(page.extract_text() or "") for page in pages)


def _clean_pdf_page(text: str) -> str:
    lines = []
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        dot_ratio = stripped.count(".") / max(len(stripped), 1)
        looks_like_toc = dot_ratio > 0.18 or "................................................................" in stripped
        looks_like_boilerplate = "GoalKicker.com" in stripped or "Free Programming Books" in stripped
        if looks_like_toc or looks_like_boilerplate:
            continue
        lines.append(stripped)
    return "\n".join(lines)


def _load_docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("Install python-docx to read DOCX files.") from exc

    doc = Document(BytesIO(data))
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)
